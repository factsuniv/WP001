#!/usr/bin/env python3
"""
Generate Word document of the Revolutionary THPU White Paper
"""

import requests
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import os

# Backend URL
BACKEND_URL = "https://9ff8b068-b843-42e4-987f-d68282246334.preview.emergentagent.com/api"

def fetch_whitepaper_data():
    """Fetch the white paper data from the backend API"""
    try:
        response = requests.get(f"{BACKEND_URL}/whitepaper", timeout=30)
        if response.status_code == 200:
            return response.json()
        else:
            print(f"Error fetching whitepaper: {response.status_code}")
            return None
    except Exception as e:
        print(f"Exception fetching whitepaper: {e}")
        return None

def add_heading_number(document, level, text):
    """Add a numbered heading to the document"""
    heading = document.add_heading(text, level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_performance_table(document):
    """Add a performance comparison table"""
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Architecture'
    header_cells[1].text = 'Energy Efficiency'
    header_cells[2].text = 'Throughput'
    header_cells[3].text = 'Latency'
    
    # Data rows
    data = [
        ['CPU', '1x', '1x', '1x'],
        ['GPU', '10x', '100x', '0.1x'],
        ['TPU', '100x', '1000x', '0.01x'],
        ['THPU', '1000x', '10000x', '0.001x']
    ]
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    return table

def generate_word_doc():
    """Generate the Word document"""
    
    # Fetch data
    print("Fetching white paper data...")
    data = fetch_whitepaper_data()
    if not data:
        print("Failed to fetch white paper data")
        return
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = data['title']
    doc.core_properties.author = ', '.join([author['name'] for author in data['authors']])
    doc.core_properties.subject = 'Revolutionary Computing Architecture'
    doc.core_properties.keywords = ', '.join(data['keywords'])
    doc.core_properties.comments = 'Generated from THPU White Paper API'
    
    # Title page
    title = doc.add_heading(data['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    doc.add_paragraph()
    
    # Authors
    for author in data['authors']:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(author['name'])
        author_run.bold = True
        author_run.font.size = Pt(14)
        
        affiliation_para = doc.add_paragraph()
        affiliation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        affiliation_run = affiliation_para.add_run(author['affiliation'])
        affiliation_run.italic = True
        affiliation_run.font.size = Pt(12)
    
    # Date
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    
    # Keywords
    doc.add_paragraph()
    keywords_para = doc.add_paragraph()
    keywords_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    keywords_run = keywords_para.add_run(f"Keywords: {', '.join(data['keywords'])}")
    keywords_run.font.size = Pt(11)
    
    # Page break
    doc.add_page_break()
    
    # Abstract
    doc.add_heading('Abstract', 1)
    abstract_para = doc.add_paragraph(data['abstract'])
    abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Key Performance Achievements
    doc.add_heading('Key Performance Achievements', 1)
    
    metrics = [
        ("Energy Efficiency Improvement", "1000x over traditional CPUs"),
        ("Throughput Increase", "100x for AI workloads"),
        ("Latency Reduction", "10x for inference tasks"),
        ("Adaptability", "Infinite through neuromorphic learning")
    ]
    
    for metric, value in metrics:
        bullet_para = doc.add_paragraph()
        bullet_para.style = 'List Bullet'
        bullet_run = bullet_para.add_run(f"{metric}: ")
        bullet_run.bold = True
        bullet_para.add_run(value)
    
    # Performance comparison table
    doc.add_heading('Performance Comparison', 2)
    add_performance_table(doc)
    
    # Page break
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    
    for i, section in enumerate(data['sections']):
        toc_para = doc.add_paragraph()
        toc_para.add_run(f"{i+1}. {section['title']}")
    
    # Page break
    doc.add_page_break()
    
    # Sections
    for i, section in enumerate(data['sections']):
        doc.add_heading(f"{i+1}. {section['title']}", 1)
        
        # Process section content
        section_content = section['content']
        paragraphs = section_content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                # Handle different formatting
                if para.startswith('**') and para.endswith('**'):
                    # Bold subheading
                    subheading = doc.add_heading(para.replace('**', '').strip(), 2)
                elif para.startswith('*') and para.endswith('*'):
                    # Italic subheading
                    subheading = doc.add_heading(para.replace('*', '').strip(), 3)
                elif para.startswith('- '):
                    # Bullet point
                    bullet_para = doc.add_paragraph()
                    bullet_para.style = 'List Bullet'
                    bullet_para.add_run(para[2:].strip())
                elif para.startswith('```') and para.endswith('```'):
                    # Code block
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run(para.replace('```', '').strip())
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(10)
                else:
                    # Regular paragraph
                    regular_para = doc.add_paragraph(para.strip())
                    regular_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add figures if any
        for figure in section.get('figures', []):
            doc.add_paragraph()
            figure_title = doc.add_paragraph()
            figure_title_run = figure_title.add_run(figure['title'])
            figure_title_run.bold = True
            figure_title_run.font.size = Pt(12)
            
            figure_caption = doc.add_paragraph(figure['caption'])
            figure_caption_run = figure_caption.runs[0]
            figure_caption_run.italic = True
            figure_caption_run.font.size = Pt(10)
    
    # References
    doc.add_page_break()
    doc.add_heading('References', 1)
    
    for i, ref in enumerate(data['references']):
        ref_para = doc.add_paragraph()
        ref_para.style = 'Normal'
        
        # Reference number
        ref_num = ref_para.add_run(f"[{i+1}] ")
        ref_num.bold = True
        
        # Reference text
        ref_text = f"{ref['title']} by {', '.join(ref['authors'])} ({ref['year']}). "
        ref_para.add_run(ref_text)
        
        # Journal name in italics
        journal_run = ref_para.add_run(ref['journal'])
        journal_run.italic = True
        
        # DOI if available
        if ref.get('doi'):
            ref_para.add_run(f". DOI: {ref['doi']}")
    
    # Save document
    filename = "/app/THPU_Revolutionary_White_Paper.docx"
    doc.save(filename)
    print(f"Word document generated successfully: {filename}")
    
    return filename

if __name__ == "__main__":
    generate_word_doc()